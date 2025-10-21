"""
Fichas_EXCEL
"""

#!/usr/bin/env python
# coding: utf-8

# In[4]:


import openai
from docx import Document
from dotenv import load_dotenv
import os

# Cargar variables de entorno desde el archivo .env
load_dotenv()

# Configura tu clave de la API de OpenAI
openai.api_key = os.getenv('OPENAI_API_KEY')

# Función para hacer una pregunta a OpenAI
def hacer_pregunta_openai(pregunta):
    try:
        respuesta = openai.chat.completions.create(
            model="gpt-4",  # Modelo más reciente compatible
            messages=[
                {"role": "system", "content": f"Eres un asistente experto en funciones excel"},
                {"role": "user", "content": pregunta}
            ],
            max_tokens=1000,  # Ajusta el número de tokens según lo necesario
            temperature=0.5
        )


        return respuesta.choices[0].message.content
    except Exception as e:
        print(f"Error al generar respuesta: {e}")
        return "No se pudo generar la respuesta"

# Función que genera la ficha para una función de Excel utilizando OpenAI
def generar_ficha_funcion(funcion):
    secciones = {
        'Nombre': funcion,
        'Descripción general': f"Describe brevemente qué hace la función de Excel {funcion} y en qué situaciones es útil.",
        'Sintaxis de la función': f"¿Cuál es la sintaxis de la función de Excel {funcion}?",
        'Explicación de los argumentos': f"Explica los argumentos que utiliza la función de Excel {funcion}.",
        'Ejemplos prácticos': f"Dame tres ejemplos prácticos del uso de la función de Excel {funcion}.",
        'Errores comunes': f"¿Cuáles son los errores más comunes al usar la función de Excel {funcion}?",
        'Aplicaciones en la práctica': f"¿En qué situaciones o áreas es más útil la función de Excel {funcion}?",
        'Notas adicionales': f"Proporciona notas adicionales o detalles importantes sobre la función de Excel {funcion}.",
        'Ejercicios propuestos': f"Proporciona un ejercicio práctico para practicar el uso de la función de Excel {funcion}."
    }

    ficha = {seccion: hacer_pregunta_openai(pregunta) if seccion != 'Nombre' else funcion for seccion, pregunta in secciones.items()}
    return ficha

# Función para guardar las fichas en un archivo Word
def generar_documento_word(lista_funciones, direccion_archivo, nombre_archivo):
    # Revisar si la ruta existe
    if not os.path.exists(direccion_archivo):
        print(f"Error: La ruta {direccion_archivo} no existe.")
        return

    doc = Document()

    # Iterar sobre cada función en la lista
    for funcion in lista_funciones:
        ficha = generar_ficha_funcion(funcion)

        if not ficha:
            print(f"Error al generar la ficha para la función: {funcion}")
            continue

        # Añadir un título con el nombre de la función
        doc.add_heading(ficha['Nombre'], level=1)

        # Añadir cada sección con su contenido
        for seccion, contenido in ficha.items():
            if seccion != 'Nombre':
                doc.add_heading(seccion, level=2)
                doc.add_paragraph(contenido)

        # Añadir un salto de página después de cada función
        doc.add_page_break()

    # Guardar el documento
    ruta_completa = os.path.join(direccion_archivo, f'{nombre_archivo}.docx')
    try:
        doc.save(ruta_completa)
        print(f"Documento '{ruta_completa}' generado con éxito.")
    except Exception as e:
        print(f"Error al guardar el documento: {e}")

# Función principal
def main():
    # Variables predefinidas
    lista_funciones =excel_functions = ['ABS', 'ACOS', 'AHORA', 'ALEATORIO.ENTRE', 'BUSCAR', 'BUSCARH', 'BUSCARV', 'BUSCARX', 'CAMBIAR', 
                                        'CARACTER', 'CELDA', 'COCIENTE', 'CODIGO', 'COEF.DE.CORREL', 'COINCIDIR', 'COLUMNA', 'COLUMNAS', 
                                        'CONCAT', 'CONCATENAR', 'CONTAR', 'CONTAR.BLANCO', 'CONTAR.SI', 'CONTAR.SI.CONJUNTO', 'CONTARA', 
                                        'COS', 'DESREF', 'DESVEST', 'DESVEST.P', 'DIA', 'DIAS', 'DIAS.LAB', 'DIAS.LAB.INTL', 'DIAS360', 
                                        'DIRECCION', 'DISTR.BINOM.N', 'DISTR.NORM', 'DISTR.NORM.ESTAND', 'DISTR.T.N', 'DURACION', 'ELEGIR', 
                                        'ENTERO', 'ES.IMPAR', 'ES.PAR', 'ESBLANCO', 'ESERROR', 'ESNUMERO', 'EXP', 'FALSO', 'FECHA', 'FECHA.MES',
                                        'FILA', 'FILAS', 'FILTRAR', 'FIN.MES', 'FORMULATEXTO', 'HIPERVINCULO', 'HORA', 'HOJAS', 'HOY', 'INDICE', 
                                        'INDIRECTO', 'INFO', 'INT.ACUM', 'INV.NORM', 'JERARQUIA', 'JERARQUIA.EQV', 'K.ESIMO.MAYOR', 'LIMPIAR', 'LN', 
                                        'LOG', 'LOG10', 'M.C.D', 'M.C.M', 'MAX', 'MAX.SI.CONJUNTO', 'MAXA', 'MAYUSC', 'MDETERM', 'MEDIA.ARMO', 
                                        'MEDIA.GEOM', 'MEDIANA', 'MES', 'MIN', 'MIN.SI.CONJUNTO', 'MINA', 'MINUSC', 'MINUTO', 'MODA', 'MODA.UNO', 
                                        'MODA.VARIOS', 'N', 'NO', 'NOMPROPIO', 'NPER', 'NUM.DE.SEMANA', 'NUMERO.ROMANO', 'O', 'ORDENAR', 
                                        'ORDENARPOR', 'PAGO', 'PENDIENTE', 'PERCENTIL', 'PI', 'POTENCIA', 'PRODUCTO', 'PROMEDIO', 'PROMEDIO.SI', 
                                        'PROMEDIO.SI.CONJUNTO', 'PROMEDIOA', 'PRONOSTICO', 'RADIANES', 'RAIZ', 'REDONDEAR', 'REDONDEAR.MAS', 
                                        'REDONDEAR.MENOS', 'REEMPLAZAR', 'REPETIR', 'RESIDUO', 'SEGUNDO', 'SENO', 'SI', 'SI.CONJUNTO', 'SI.ERROR', 
                                        'SIFECHA', 'SIGNO', 'SUBTOTALES', 'SUMA', 'SUMA.CUADRADOS', 'SUMA.SERIES', 'SUMAPRODUCTO', 'SUMAR.SI', 
                                        'SUMAR.SI.CONJUNTO', 'SUSTITUIR', 'T', 'TAN', 'TASA', 'TENDENCIA', 'TEXTO', 'TRANSPONER', 'TRUNCAR', 
                                        'UNICOS', 'VALOR', 'VALOR.NUMERO', 'VAR.P', 'VAR.S', 'VERDADERO', 'VF', 'VNA', 'Y', 'XOR']

    nombre_archivo = "Fichas_Funciones_Excel_OpenAI"
    direccion_archivo = f'C:\\Users\\HP\\Desktop\\LIBROS PERSO\\EXCEL FUNCIONES'

    # Generar el documento
    generar_documento_word(lista_funciones, direccion_archivo, nombre_archivo)

if __name__ == "__main__":
    main()


# In[ ]:




